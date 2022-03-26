#!/usr/bin/env python
# coding: utf-8

import re

from io import BytesIO
from urllib import request

import yaml
import docx


def clean_line(long_line):
    lines = []
    
    for line in long_line.split('\n'):
        line = line.strip()
        line = line.strip('\uFEFF')
        lines.append(line)

    return lines


ts_two_digit_pattern = re.compile(r'^\[?(\d+):(\d+)\]?$')
ts_three_digit_pattern = re.compile(r'^\[?(\d+):(\d+):(\d+)\]?$')


def try_parse_time(line):
    match = ts_two_digit_pattern.match(line)

    if match is not None:
        m, s = int(match.group(1)), int(match.group(2))
        h = m // 60
        m = m % 60
        return h, m, s

    match = ts_three_digit_pattern.match(line)
    if match is not None:
        h = int(match.group(1))
        m = int(match.group(2))
        s = int(match.group(3))
        return h, m, s

    return None


def format_time(h, m, s):
    if h > 0:
        return '%d:%02d:%02d' % (h, m, s)
    return '%d:%02d' % (m, s)


def url_as_bytes_stream(url):
    with request.urlopen(url) as resp:
        buffer = resp.read()
    stream = BytesIO(buffer)
    return stream


def parse_doc(doc):
    h = None
    m = None
    s = None
    total_sec = None
    speaker = None
    after_time = False

    elements = []

    for p in doc.paragraphs:
        style = p.style.name.lower()
        lines = clean_line(p.text)

        for line in lines:
            if len(line) == 0:
                continue

            if style.startswith('heading'):
                element = {
                    'header': line
                }
                elements.append(element)
                continue

            assert style == 'normal'

            if after_time == True:
                after_time = False
                speaker = line
                continue

            maybe_time = try_parse_time(line)

            if maybe_time is not None:
                h, m, s = maybe_time
                total_sec = h * 60 * 60 + m * 60 + s
                after_time = True
                continue

            assert speaker is not None
            assert total_sec is not None

            element = {
                'time': format_time(h, m, s),
                'sec': total_sec,
                'who': speaker,
                'line': line
            }

            elements.append(element)

    return elements


def read_docx(input_file):
    if input_file.startswith('http'):
        stream = url_as_bytes_stream(input_file)
        doc = docx.Document(stream)
    else:
        doc = docx.Document(input_file)

    return parse_doc(doc)


def read_docx_base64(input_file):
    pass


def extract_timecodes(elements):
    timecodes = [
        '00:00 DataTalks.Club intro'
    ]

    for i in range(len(elements)):
        row = elements[i]
        if 'header' not in row:
            continue

        next_row = elements[i + 1]

        time = next_row['time']

        if len(time) < 5:
            time = '0' + time

        tc = '%s %s' % (time, row['header'])
        timecodes.append(tc)

    return timecodes


def process_trascript(input_file, output_file_timecodes,
        output_file_transcript):
    elements = read_docx(input_file)

    print(f'writing transcript to {output_file_transcript}...')
    with open(output_file_transcript, 'w', encoding='utf-8') as f_out:
        yaml.dump({'transcript': elements}, f_out)

    timecodes = extract_timecodes(elements)

    print(f'writing timecodes to {output_file_timecodes}...')
    with open(output_file_timecodes, 'w', encoding='utf-8') as f_out:
        for line in timecodes:
            f_out.write(line)
            f_out.write('\n')
